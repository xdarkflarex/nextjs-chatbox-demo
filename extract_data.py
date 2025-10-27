"""
Script trích xuất nội dung từ Word/Excel sang JSON cho RAG
Cần cài: pip install python-docx openpyxl
"""

import json
import os
from pathlib import Path

try:
    from docx import Document
    from openpyxl import load_workbook
except ImportError:
    print("❌ Cần cài thư viện:")
    print("pip install python-docx openpyxl")
    exit(1)

def extract_docx(file_path):
    """Trích xuất text từ file Word"""
    doc = Document(file_path)
    content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
    
    # Trích xuất từ tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                content.append(" | ".join(row_data))
    
    return "\n".join(content)

def extract_xlsx(file_path):
    """Trích xuất data từ file Excel"""
    wb = load_workbook(file_path, data_only=True)
    content = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        content.append(f"\n=== {sheet_name} ===\n")
        
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            row_text = " | ".join([c for c in row_data if c.strip()])
            if row_text.strip():
                content.append(row_text)
    
    return "\n".join(content)

def main():
    data_update_dir = Path("data update")
    output_dir = Path("app/public/data")
    
    if not data_update_dir.exists():
        print(f"❌ Không tìm thấy thư mục: {data_update_dir}")
        return
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Danh sách file cần trích xuất
    files = list(data_update_dir.glob("*.docx")) + list(data_update_dir.glob("*.xlsx"))
    
    if not files:
        print("❌ Không tìm thấy file Word/Excel trong thư mục 'data update'")
        return
    
    print(f"📂 Tìm thấy {len(files)} file\n")
    
    all_data = []
    
    for file_path in files:
        print(f"📄 Đang xử lý: {file_path.name}")
        
        try:
            if file_path.suffix == ".docx":
                content = extract_docx(file_path)
            elif file_path.suffix == ".xlsx":
                content = extract_xlsx(file_path)
            else:
                continue
            
            # Tạo entry cho RAG
            entry = {
                "source": file_path.name,
                "type": "document",
                "content": content,
                "metadata": {
                    "filename": file_path.name,
                    "extension": file_path.suffix
                }
            }
            
            all_data.append(entry)
            
            # Lưu file riêng
            output_file = output_dir / f"{file_path.stem}.txt"
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(content)
            
            print(f"   ✅ Đã lưu: {output_file.name}")
            print(f"   📊 Độ dài: {len(content)} ký tự\n")
            
        except Exception as e:
            print(f"   ❌ Lỗi: {e}\n")
    
    # Lưu tất cả vào 1 file JSON
    output_json = output_dir / "extracted_data.json"
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ Hoàn thành!")
    print(f"📁 Đã lưu {len(all_data)} file vào: {output_dir}")
    print(f"📄 File tổng hợp: {output_json}")
    
    # Tạo file RAG format
    print("\n📝 Tạo file RAG format...")
    rag_entries = []
    
    for idx, data in enumerate(all_data, 1):
        # Split content thành các đoạn nhỏ
        lines = data["content"].split("\n")
        current_chunk = []
        
        for line in lines:
            if line.strip():
                current_chunk.append(line)
                
                # Mỗi chunk ~500 ký tự
                if len("\n".join(current_chunk)) > 500:
                    rag_entry = {
                        "id": f"doc_{idx}_{len(rag_entries)}",
                        "question": "",  # Để trống, sẽ dùng cho semantic search
                        "answer": "\n".join(current_chunk),
                        "category": "document",
                        "keywords": [data["source"]],
                        "source": data["source"]
                    }
                    rag_entries.append(rag_entry)
                    current_chunk = []
        
        # Chunk cuối
        if current_chunk:
            rag_entry = {
                "id": f"doc_{idx}_{len(rag_entries)}",
                "question": "",
                "answer": "\n".join(current_chunk),
                "category": "document",
                "keywords": [data["source"]],
                "source": data["source"]
            }
            rag_entries.append(rag_entry)
    
    # Lưu RAG format
    rag_output = output_dir / "rag_documents.json"
    with open(rag_output, "w", encoding="utf-8") as f:
        json.dump(rag_entries, f, ensure_ascii=False, indent=2)
    
    print(f"✅ Đã tạo {len(rag_entries)} RAG entries")
    print(f"📄 File RAG: {rag_output}")
    
    print("\n" + "="*50)
    print("🎯 HƯỚNG DẪN SỬ DỤNG:")
    print("="*50)
    print("1. File text đã trích xuất: app/public/data/*.txt")
    print("2. File JSON tổng hợp: app/public/data/extracted_data.json")
    print("3. File RAG format: app/public/data/rag_documents.json")
    print("\n4. Để sử dụng trong API:")
    print("   - Import rag_documents.json vào RAG system")
    print("   - Hoặc merge với rag_all.json hiện có")
    print("\n5. Test:")
    print("   python test_rag.py")

if __name__ == "__main__":
    main()
